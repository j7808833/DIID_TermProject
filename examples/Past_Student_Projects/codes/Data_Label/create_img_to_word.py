import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os
import glob

# === 1. 設定 ===
excel_folder = 'D:/AI/excel_preview/0606/'  # 放 Excel 的資料夾
img_root_folder = 'D:/AI/excel_preview/gy_charts_0606'
#word_output = 'D:/AI/excel_previewgY_charts_report_0606.docx'

# === 2. 讀取所有 Excel 檔並合併 ===
all_data = []
for file_path in glob.glob(os.path.join(excel_folder, '*.xlsx')):
    df = pd.read_excel(file_path)
    df['source_file'] = os.path.basename(file_path)  # 加來源檔案名
    all_data.append(df)
    print("file_path="+file_path)

df = pd.concat(all_data, ignore_index=True)

# === 3. 建立圖片資料夾 ===
os.makedirs(img_root_folder, exist_ok=True)

# === 4. 繪製並儲存每段 gY 折線圖 ===
for sample_idx, group in df.groupby("sample_index"):
    label = group['label'].iloc[0] or 'unknown'
    subfolder = os.path.join(img_root_folder, label)
    os.makedirs(subfolder, exist_ok=True)

    # 若有來源檔名，加入圖名以示區別
    source_file = group['source_file'].iloc[0].replace('.xlsx', '')
    filename = f"gY_sample_{int(sample_idx)}_{label}_{source_file}.png"
    filepath = os.path.join(subfolder, filename)

    plt.figure()
    plt.plot(group['gY'])
    plt.title(f"Sample {int(sample_idx)} - gY ({label})\n{source_file}")
    plt.xlabel("Index")
    plt.ylabel("gY")
    plt.grid(True)
    plt.savefig(filepath)
    plt.close()
print(f"文件儲存完成")
'''
# === 5. 建立 Word 文件 ===
doc = Document()
doc.add_heading('gY 折線圖報告（多檔案支援 / 依 label 分類）', level=1)

for label in sorted(os.listdir(img_root_folder)):
    subfolder = os.path.join(img_root_folder, label)
    if not os.path.isdir(subfolder):
        continue

    doc.add_heading(f"Label: {label}", level=2)
    image_files = sorted(
        glob.glob(os.path.join(subfolder, f"gY_sample_*_{label}_*.png")),
        key=lambda x: (
            int(os.path.basename(x).split('_')[2]),  # sample_index
            os.path.basename(x)                      # 次排序：檔名
        )
    )

    for img_path in image_files:
        doc.add_paragraph(os.path.basename(img_path), style='Heading 3')
        doc.add_picture(img_path, width=Inches(5.5))

doc.save(word_output)
print(f"Word 文件儲存完成：{word_output}")
'''